from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Kategori Test Laboratorium', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Berikut adalah pengelompokan daftar test laboratorium dari sistem (duplikat telah digabungkan agar lebih rapi):")
    
    categories = {
        "1. Urinalysis Dipstick": [
            "Urine Dipstick (Urine Dip)"
        ],
        "2. Serology & Rapid Test": [
            "ABO & Rh(D) (Blood Group)",
            "HIV 1&2 (HIV)",
            "H.pylori",
            "Malaria Ag (Malaria)",
            "Dengue NS1",
            "Dengue IgG/IgM (Dengue)",
            "VDRL / RPR / Syphilis",
            "TPHA",
            "Widal",
            "HBsAg (Hep. B)",
            "Hep. C (HCV)",
            "HCG (Pregnancy Test)",
            "Gonorrhea",
            "Prenatal Panel"
        ],
        "3. Clinical Biochemistry": [
            "BUN (Blood Urea Nitrogen)",
            "CRE (Creatinine)",
            "GOT (AST / SGOT)",
            "GPT (ALT / SGPT)",
            "ALP (Alkaline Phosphatase)",
            "UA (Uric Acid)",
            "TCHO (Total Cholesterol)",
            "TBIL (Bilirubin Total)",
            "DBIL (Bilirubin Direct)",
            "GLU (Glucose) / B.S (Blood Sugar)",
            "LFT (Liver Function Test)",
            "U&E (Urea & Electrolytes)"
        ],
        "4. Haematology & Immunology": [
            "CBC (Complete Blood Count)",
            "Hb (Hemoglobin)",
            "INR (International Normalized Ratio)",
            "Smear Morphology"
        ],
        "5. Micro & Parasitology Examination": [
            "G.O. Smear",
            "Urine Sediment",
            "Gram Stain",
            "Stool O&P"
        ]
    }
    
    for cat_name, tests in categories.items():
        doc.add_heading(cat_name, level=1)
        for test in tests:
            doc.add_paragraph(test, style='List Bullet')
            
    doc.save(r'D:\pasient_system\Kategori_Test_Lab.docx')
    print("Document saved successfully.")

if __name__ == '__main__':
    create_doc()
